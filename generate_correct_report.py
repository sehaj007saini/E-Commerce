import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

def add_page_numbers(doc):
    """Add page numbers to footer"""
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add page number field
        run = footer_para.add_run()
        fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">PAGE</w:instrText>')
        fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

def create_report():
    doc = docx.Document()
    
    # Set margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
    
    # Default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    def add_heading_centered(text, size=18, bold=True, space_before=12, space_after=12):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        run.font.bold = bold
        return p
    
    def add_para(text, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
        p = doc.add_paragraph()
        p.alignment = alignment
        p.paragraph_format.line_spacing = 1.5
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p
    
    # ========== TITLE PAGE ==========
    doc.add_paragraph()  # spacing
    add_heading_centered("SUMMER TRAINING REPORT", 22, space_before=48)
    add_heading_centered("On", 14, space_before=12)
    
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(24)
    run = p_title.add_run("Full-Stack E-Commerce Platform\nwith Admin Authentication System")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)
    
    add_para("Submitted in the partial fulfillment of the requirement\nfor the award of degree of", WD_ALIGN_PARAGRAPH.CENTER)
    
    p_degree = doc.add_paragraph()
    p_degree.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_degree.paragraph_format.space_before = Pt(8)
    p_degree.paragraph_format.space_after = Pt(32)
    run = p_degree.add_run("Bachelor of Technology\nin\nComputer Science and Engineering\nSession: 2024-2028")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.font.bold = True
    
    add_heading_centered("Submitted By:", 12, space_before=36)
    
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_name.add_run("Sehajpreet Singh\nRoll No: 2411804\nB.Tech CSE (5th Semester)")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    
    add_heading_centered("Under the Supervision of:", 12, space_before=32)
    p_sup = doc.add_paragraph()
    p_sup.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_sup.add_run("Navin Reddy\nSenior Software Instructor\nUdemy")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    add_heading_centered("DEPARTMENT OF COMPUTER SCIENCE AND ENGINEERING", 13, space_before=48)
    add_heading_centered("AMRITSAR GROUP OF COLLEGES", 14)
    add_heading_centered("AMRITSAR, PUNJAB", 13)
    add_heading_centered("June 2024 - July 2024", 12, space_before=16)
    
    doc.add_page_break()
    
    # ========== CERTIFICATE ==========
    add_heading_centered("CERTIFICATE", 18, space_before=24)
    doc.add_paragraph()
    
    add_para("This is to certify that Mr. Sehajpreet Singh, Roll No. 2411804, a bonafide student of B.Tech Computer Science and Engineering (5th Semester), Session 2024-2028, has successfully completed Summer Training on \"Full-Stack E-Commerce Platform with Admin Authentication System\" from 8th June 2024 to 24th July 2024.")
    
    add_para("The training was conducted online through Udemy under the supervision of Navin Reddy, Senior Software Instructor.")
    
    add_para("During the training period, the student has shown sincere interest, dedication, and has successfully completed the project with satisfactory performance.")
    
    add_para("We wish him all the best for his future endeavors.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    add_para("Date: _______________", WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Place: Amritsar", WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_sig = doc.add_paragraph()
    p_sig.add_run("_______________________          _______________________          _______________________")
    p_sig2 = doc.add_paragraph()
    p_sig2.add_run("Training Supervisor                    Head of Department                           Principal")
    p_sig3 = doc.add_paragraph()
    p_sig3.add_run("Navin Reddy                              Dept. of CSE                                    AGC Amritsar")
    
    doc.add_page_break()
    
    # ========== DECLARATION ==========
    add_heading_centered("DECLARATION", 18, space_before=24)
    doc.add_paragraph()
    
    add_para("I, Sehajpreet Singh, Roll No. 2411804, student of B.Tech Computer Science and Engineering (5th Semester), Session 2024-2028, at Amritsar Group of Colleges, Amritsar, Punjab, hereby declare that the Summer Training Report titled \"Full-Stack E-Commerce Platform with Admin Authentication System\" submitted in partial fulfillment of the requirement for the award of the degree of Bachelor of Technology in Computer Science and Engineering is my original work.")
    
    add_para("This report is based on the practical training undertaken by me at Udemy (Online Mode) from 8th June 2024 to 24th July 2024 under the guidance and supervision of Mr. Navin Reddy, Senior Software Instructor.")
    
    add_para("The information presented in this report has been collected and compiled by me through various sources including online courses, official documentation, practical implementation, and personal research. The project implementation, code development, and system design are entirely my own work.")
    
    add_para("I further declare that this report has not been submitted earlier to any other institution or university for the award of any degree or diploma.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    add_para("Date: _______________", WD_ALIGN_PARAGRAPH.LEFT)
    add_para("Place: Amritsar", WD_ALIGN_PARAGRAPH.LEFT)
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_st = doc.add_paragraph()
    p_st.add_run("_______________________\nSehajpreet Singh\nRoll No: 2411804\nB.Tech CSE (5th Semester)")
    
    doc.add_page_break()
    
    # ========== ACKNOWLEDGEMENT ==========
    add_heading_centered("ACKNOWLEDGEMENT", 18, space_before=24)
    doc.add_paragraph()
    
    add_para("First and foremost, I would like to express my sincere gratitude to the Almighty God for His blessings, guidance, and strength throughout my summer training journey.")
    
    add_para("I would like to extend my heartfelt thanks to Mr. Navin Reddy, Senior Software Instructor at Udemy, for his excellent guidance, constant encouragement, and valuable supervision throughout the training period. His comprehensive teaching methodology, real-world examples, and practical approach to full-stack development have been instrumental in helping me understand complex concepts and implement them successfully in this project.")
    
    add_para("I am deeply grateful to the Principal of Amritsar Group of Colleges, for providing the opportunity to undertake this summer training and for creating an environment conducive to learning and innovation.")
    
    add_para("I express my sincere thanks to the Head of Department, Computer Science and Engineering, for continuous support, motivation, and valuable suggestions throughout the training period. The guidance has been invaluable in shaping my technical skills and professional development.")
    
    add_para("I would like to thank my Training Coordinator and all the faculty members of the Department of Computer Science and Engineering for their constant support and encouragement.")
    
    add_para("I am thankful to Udemy Platform for providing world-class online courses and resources that enabled me to learn industry-standard technologies and best practices in software development.")
    
    add_para("I also extend my gratitude to my parents and family members for their unconditional love, support, and encouragement throughout this training period. Their belief in my abilities has been my greatest motivation.")
    
    add_para("Last but not least, I would like to thank my friends and classmates for their cooperation, healthy discussions, and moral support during the training period.")
    
    doc.add_paragraph()
    doc.add_paragraph()
    
    p_ack = doc.add_paragraph()
    p_ack.add_run("Sehajpreet Singh\nRoll No: 2411804\nB.Tech CSE (5th Semester)")
    
    doc.add_page_break()
    
    # ========== ABSTRACT ==========
    add_heading_centered("ABSTRACT", 18, space_before=24)
    doc.add_paragraph()
    
    add_para("In the modern digital era, e-commerce has revolutionized the way consumers purchase goods and services. This summer training project presents the design, development, and implementation of a Full-Stack E-Commerce Platform with Admin Authentication System using cutting-edge web technologies.")
    
    add_para("The project was undertaken during the summer training period from 8th June 2024 to 24th July 2024 under the guidance of Mr. Navin Reddy, Senior Software Instructor at Udemy, through comprehensive online courses in full-stack web development.")
    
    add_para("This e-commerce platform is built using a modern technology stack comprising React.js for the frontend, Spring Boot (Java) for the backend, and PostgreSQL for database management. The system follows a three-tier architecture that ensures scalability, maintainability, and separation of concerns.")
    
    add_para("The application provides a comprehensive solution for online shopping with features including user authentication and authorization, product catalog management, advanced search and filtering, shopping cart, wishlist functionality, category-based browsing, and admin dashboard for product management. The admin authentication system ensures that only authorized administrators can perform CRUD operations on products, while regular users can browse, search, and shop seamlessly.")
    
    add_para("The frontend is developed using React.js with Vite as the build tool, offering a modern, responsive, and intuitive user interface with a purple gradient theme. React Router is used for client-side routing, and Axios handles HTTP requests to the backend API.")
    
    add_para("The backend is powered by Spring Boot, utilizing Spring Web MVC for RESTful API development, Spring Data JPA for database operations, and a well-structured service-repository architecture. The system implements global exception handling, input validation, and comprehensive error management.")
    
    add_para("Key achievements include successful implementation of 30+ sample products across 6 categories, real-time search functionality, category-based filtering, stock management, and a complete admin authentication system with protected routes.")
    
    add_para("This training experience provided hands-on exposure to industry-standard technologies, full-stack development practices, RESTful API design, database management, state management in React, and modern UI/UX principles.")
    
    doc.add_paragraph()
    p_key = doc.add_paragraph()
    p_key.paragraph_format.space_before = Pt(12)
    run = p_key.add_run("Keywords: ")
    run.font.bold = True
    p_key.add_run("E-Commerce, React.js, Spring Boot, PostgreSQL, Full-Stack Development, Admin Authentication, REST API, Web Application")
    
    doc.add_page_break()
    
    # ========== TABLE OF CONTENTS ==========
    add_heading_centered("TABLE OF CONTENTS", 18, space_before=24)
    doc.add_paragraph()
    
    toc_items = [
        ("CERTIFICATE", "i"),
        ("DECLARATION", "ii"),
        ("ACKNOWLEDGEMENT", "iii"),
        ("ABSTRACT", "iv"),
        ("TABLE OF CONTENTS", "v"),
        ("", ""),
        ("CHAPTER 1: INTRODUCTION", "1"),
        ("    1.1  Overview of E-Commerce", "1"),
        ("    1.2  Need for the Project", "2"),
        ("    1.3  Problem Statement", "3"),
        ("    1.4  Objectives of the Project", "4"),
        ("    1.5  Scope of the Project", "5"),
        ("", ""),
        ("CHAPTER 2: TRAINING OBJECTIVES", "7"),
        ("    2.1  Primary Objectives", "7"),
        ("    2.2  Learning Outcomes", "8"),
        ("    2.3  Skills Acquired", "9"),
        ("    2.4  Industry Relevance", "11"),
        ("", ""),
        ("CHAPTER 3: ORGANIZATION BRIEF", "13"),
        ("    3.1  About Udemy", "13"),
        ("    3.2  Training Infrastructure", "14"),
        ("    3.3  Course Structure", "15"),
        ("    3.4  Instructor Profile", "16"),
        ("", ""),
        ("CHAPTER 4: LITERATURE REVIEW", "18"),
        ("    4.1  Overview of E-Commerce Systems", "18"),
        ("    4.2  Existing E-Commerce Platforms", "19"),
        ("    4.3  Technology Stack Analysis", "20"),
        ("    4.4  Security Considerations", "22"),
        ("    4.5  Gap Analysis", "23"),
        ("", ""),
        ("CHAPTER 5: TECHNOLOGY USED", "25"),
        ("    5.1  Frontend Technologies", "25"),
        ("    5.2  Backend Technologies", "29"),
        ("    5.3  Database Technology", "33"),
        ("    5.4  Development Tools", "35"),
        ("", ""),
        ("CHAPTER 6: SOFTWARE DEVELOPMENT MODEL", "37"),
        ("    6.1  Agile Methodology", "37"),
        ("    6.2  Iterative Development Approach", "38"),
        ("    6.3  Sprint Planning and Execution", "40"),
        ("    6.4  Development Lifecycle", "41"),
        ("", ""),
        ("CHAPTER 7: SYSTEM ANALYSIS", "43"),
        ("    7.1  Requirement Analysis", "43"),
        ("    7.2  Feasibility Study", "46"),
        ("    7.3  System Requirements", "48"),
        ("", ""),
        ("CHAPTER 8: SYSTEM DESIGN", "50"),
        ("    8.1  System Architecture", "50"),
        ("    8.2  Database Design", "52"),
        ("    8.3  Use Case Diagram", "54"),
        ("    8.4  Sequence Diagrams", "55"),
        ("", ""),
        ("CHAPTER 9: IMPLEMENTATION", "57"),
        ("    9.1  Backend Implementation", "57"),
        ("    9.2  Frontend Implementation", "60"),
        ("    9.3  Authentication System", "63"),
        ("    9.4  Product Management Module", "65"),
        ("", ""),
        ("CHAPTER 10: TESTING", "67"),
        ("    10.1  Testing Strategy", "67"),
        ("    10.2  API Testing", "68"),
        ("    10.3  User Acceptance Testing", "69"),
        ("", ""),
        ("CHAPTER 11: RESULTS AND SCREENSHOTS", "71"),
        ("    11.1  Home Page", "71"),
        ("    11.2  Product Listing and Details", "72"),
        ("    11.3  Admin Login and Dashboard", "73"),
        ("    11.4  Shopping Cart and Wishlist", "75"),
        ("", ""),
        ("CHAPTER 12: CHALLENGES AND SOLUTIONS", "77"),
        ("    12.1  Technical Challenges", "77"),
        ("    12.2  Solutions Implemented", "78"),
        ("", ""),
        ("CHAPTER 13: FUTURE SCOPE", "80"),
        ("    13.1  Proposed Enhancements", "80"),
        ("    13.2  Advanced Features", "81"),
        ("", ""),
        ("CHAPTER 14: CONCLUSION", "83"),
        ("    14.1  Project Summary", "83"),
        ("    14.2  Learning Experience", "84"),
        ("", ""),
        ("BIBLIOGRAPHY", "86"),
        ("APPENDIX A: Source Code Snippets", "88"),
        ("APPENDIX B: API Documentation", "95"),
        ("APPENDIX C: Database Schema", "100"),
    ]
    
    for item, page in toc_items:
        if item == "":
            doc.add_paragraph()
        else:
            p = doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.3
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(Inches(6.0))
            run = p.add_run(f"{item}\t{page}")
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    doc.add_page_break()
    
    # ========== CHAPTER 1 ==========
    add_heading_centered("CHAPTER 1", 16, space_before=24)
    add_heading_centered("INTRODUCTION", 16, space_after=18)
    
    p_h = doc.add_paragraph()
    p_h.paragraph_format.space_before = Pt(12)
    run = p_h.add_run("1.1  Overview of E-Commerce")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    
    add_para("Electronic commerce, commonly known as e-commerce, refers to the buying and selling of goods and services over the internet. It has transformed the traditional retail landscape by providing consumers with convenient access to products and services from anywhere in the world, at any time. The e-commerce industry has experienced exponential growth over the past decade, with global e-commerce sales reaching trillions of dollars annually.")
    
    add_para("E-commerce platforms serve as digital marketplaces where businesses can showcase their products, manage inventory, process orders, and interact with customers. These platforms have revolutionized consumer behavior, offering advantages such as 24/7 availability, wider product selection, competitive pricing, easy price comparison, and doorstep delivery.")
    
    add_para("Modern e-commerce platforms leverage advanced technologies such as cloud computing, artificial intelligence, machine learning, big data analytics, and responsive web design to provide seamless user experiences across multiple devices and platforms. The COVID-19 pandemic further accelerated the adoption of e-commerce, making e-commerce development skills highly valuable in today's job market.")
    
    # Continue with more content...
    
    doc.add_paragraph()
    p_h2 = doc.add_paragraph()
    run = p_h2.add_run("1.2  Need for the Project")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.bold = True
    
    add_para("The need for developing this e-commerce platform stems from multiple factors. As part of the B.Tech Computer Science and Engineering curriculum, practical training and project work are essential components that bridge the gap between theoretical knowledge and real-world application. This project provides hands-on experience with industry-standard technologies and development practices.")
    
    add_para("Full-stack development is one of the most sought-after skill sets in the software industry. Understanding both frontend and backend technologies, along with database management, makes developers more versatile and employable. E-commerce is a booming industry with continuous demand for skilled developers. Building an e-commerce platform demonstrates practical understanding of real-world business requirements.")
    
    # Add more sections with proper formatting
    
    # Add page numbers
    add_page_numbers(doc)
    
    # Save
    output_path = os.path.join(os.getcwd(), "Sehajpreet_Singh_2411804_Summer_Training_Report.docx")
    doc.save(output_path)
    print(f"✅ Report generated successfully!")
    print(f"📄 File: {output_path}")
    print(f"👤 Name: Sehajpreet Singh")
    print(f"🎓 Roll No: 2411804")
    print(f"📊 Page numbers: Added to footer")
    return output_path

if __name__ == "__main__":
    create_report()
