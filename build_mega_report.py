import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

def set_cell_margins(cell, top=120, bottom=120, left=160, right=160):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>\n'
        f'  <w:top w:w="{top}" w:type="dxa"/>\n'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>\n'
        f'  <w:left w:w="{left}" w:type="dxa"/>\n'
        f'  <w:right w:w="{right}" w:type="dxa"/>\n'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_table_borders(table, color="B0C4DE", sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_mega_report():
    doc = docx.Document()
    
    # 1.0 inch margins on all sides
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base Typography Settings
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_font.color.rgb = RGBColor(0x11, 0x11, 0x11)
    normal_style.paragraph_format.line_spacing = 1.5
    normal_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal_style.paragraph_format.space_after = Pt(6)

    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(14)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x20, 0x60)
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
        return p

    def add_p(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        return p

    def add_bullet(text, bold_prefix=""):
        p = doc.add_paragraph(style='List Bullet')
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(4)
        if bold_prefix:
            r_bold = p.add_run(bold_prefix)
            r_bold.font.name = 'Times New Roman'
            r_bold.font.size = Pt(12)
            r_bold.font.bold = True
        r_text = p.add_run(text)
        r_text.font.name = 'Times New Roman'
        r_text.font.size = Pt(12)
        return p

    def add_code_block(code_str):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        set_cell_shading(cell, "F4F6F8")
        set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
        
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(0)
        cp.paragraph_format.line_spacing = 1.15
        r = cp.add_run(code_str)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_table_custom(headers, data, col_widths=None, caption=""):
        table = doc.add_table(rows=len(data) + 1, cols=len(headers))
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(table, color="B0C4DE", sz="4")

        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_shading(hdr_cells[i], "1F497D")
            set_cell_margins(hdr_cells[i], top=120, bottom=120, left=140, right=140)
            p = hdr_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(0)
            for run in p.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

        for r_idx, row_data in enumerate(data):
            row_cells = table.rows[r_idx + 1].cells
            shd_color = "F9FAFB" if r_idx % 2 == 1 else "FFFFFF"
            for c_idx, cell_value in enumerate(row_data):
                row_cells[c_idx].text = str(cell_value)
                set_cell_shading(row_cells[c_idx], shd_color)
                set_cell_margins(row_cells[c_idx], top=100, bottom=100, left=140, right=140)
                p = row_cells[c_idx].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                if c_idx == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        if col_widths:
            for row in table.rows:
                for i, w in enumerate(col_widths):
                    row.cells[i].width = Inches(w)

        if caption:
            cp = doc.add_paragraph()
            cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cp.paragraph_format.space_before = Pt(4)
            cp.paragraph_format.space_after = Pt(12)
            c_run = cp.add_run(caption)
            c_run.font.name = 'Times New Roman'
            c_run.font.size = Pt(10)
            c_run.font.italic = True
            c_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    # -------------------------------------------------------------
    # 1. TITLE PAGE (Appendix A Layout)
    # -------------------------------------------------------------
    p_college = doc.add_paragraph()
    p_college.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_college.paragraph_format.space_before = Pt(24)
    p_college.paragraph_format.space_after = Pt(18)
    r_col = p_college.add_run("AMRITSAR GROUP OF COLLEGES")
    r_col.font.name = 'Times New Roman'
    r_col.font.size = Pt(22)
    r_col.font.bold = True
    r_col.font.color.rgb = RGBColor(0x00, 0x20, 0x60)

    p_rep = doc.add_paragraph()
    p_rep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_rep.paragraph_format.space_after = Pt(8)
    r_rep = p_rep.add_run("Summer Training Report\nOn")
    r_rep.font.name = 'Times New Roman'
    r_rep.font.size = Pt(14)
    r_rep.font.bold = True

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(18)
    r_t = p_title.add_run("SpringEcom: Enterprise AI-Powered Full-Stack E-Commerce Platform")
    r_t.font.name = 'Times New Roman'
    r_t.font.size = Pt(18)
    r_t.font.bold = True
    r_t.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    r_s = p_sub.add_run("Submitted in partial fulfillment of the requirement for the award of degree of\n")
    r_s.font.name = 'Times New Roman'
    r_s.font.size = Pt(12)
    r_s_b = p_sub.add_run("Bachelor of Technology\nIn\nCOMPUTER SCIENCE & ENGINEERING\nBatch (2024-2028)")
    r_s_b.font.name = 'Times New Roman'
    r_s_b.font.size = Pt(14)
    r_s_b.font.bold = True

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(36)

    sub_table = doc.add_table(rows=2, cols=2)
    sub_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c00 = sub_table.cell(0, 0)
    c01 = sub_table.cell(0, 1)
    c10 = sub_table.cell(1, 0)
    c11 = sub_table.cell(1, 1)

    c00.paragraphs[0].add_run("Submitted to").font.bold = True
    c01.paragraphs[0].add_run("Submitted by").font.bold = True
    c10.paragraphs[0].add_run("Department of CSE\nAmritsar Group of Colleges")
    c11.paragraphs[0].add_run("Sehajpreet Singh (Roll No. 2411804)")

    for row in sub_table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.3
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)

    p_dept = doc.add_paragraph()
    p_dept.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_dept.paragraph_format.space_before = Pt(48)
    r_d = p_dept.add_run("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING\nAMRITSAR GROUP OF COLLEGES, AMRITSAR\n2026")
    r_d.font.name = 'Times New Roman'
    r_d.font.size = Pt(13)
    r_d.font.bold = True

    doc.add_page_break()

    # -------------------------------------------------------------
    # PRELIMINARY PAGES
    # -------------------------------------------------------------
    add_title("SUMMER TRAINING CERTIFICATE")
    add_p("This is to certify that Mr. Sehajpreet Singh, Roll No. 2411804, a student of Bachelor of Technology in Computer Science & Engineering at Amritsar Group of Colleges, has successfully completed Summer Training from 8th June 2024 to 24th July 2024.")
    add_p("During this period, they developed an enterprise-grade full-stack project titled \"SpringEcom: Enterprise AI-Powered Full-Stack E-Commerce Platform\" using Spring Boot 3, Java 17/21, Spring Data JPA, Spring Security (JWT), React 18, Vite, and Google Gemini Generative AI API.")
    add_p("Their performance during the training was found to be EXCELLENT. We wish them all success in their future academic and professional endeavors.")
    
    p_sig = doc.add_paragraph()
    p_sig.paragraph_format.space_before = Pt(60)
    p_sig.add_run("_________________________\t\t\t_________________________\nTraining Supervisor / HOD\t\t\tExternal Examiner\nDept. of Computer Science & Engineering\t\tAmritsar Group of Colleges")

    doc.add_page_break()

    add_title("DECLARATION")
    add_p("We hereby declare that the project report entitled \"SpringEcom: Enterprise AI-Powered Full-Stack E-Commerce Platform\" submitted in partial fulfillment of the requirements for the award of the degree of Bachelor of Technology in Computer Science & Engineering to Amritsar Group of Colleges, Amritsar, is an authentic record of our own work carried out during the period from 4th June 2026 to 10th July 2026 under the supervision of the Department of Computer Science & Engineering.")
    add_p("The matter contained in this report has not been submitted by us for the award of any other degree or diploma to any other institute or university. All sources of knowledge, libraries, framework documentations, and third-party APIs utilized in this software development have been duly acknowledged.")
    
    p_dec_sig = doc.add_paragraph()
    p_dec_sig.paragraph_format.space_before = Pt(80)
    p_dec_sig.add_run("Sehajpreet Singh (Roll No. 2411804)\nBranch: Computer Science & Engineering\nAmritsar Group of Colleges")

    doc.add_page_break()

    add_title("ACKNOWLEDGEMENT")
    add_p("We would like to express our deepest gratitude and sincere appreciation to the Department of Computer Science & Engineering, Amritsar Group of Colleges, for providing us with the opportunity, guidance, and state-of-the-art laboratory infrastructure to complete our Summer Training (AGCS-21510).")
    add_p("We extend our heartfelt thanks to our Training Supervisors and Faculty Members for their continuous support, valuable insights, constructive feedback, and constant encouragement throughout the conceptualization, architectural design, coding, testing, and deployment phases of the SpringEcom project.")
    add_p("Special thanks are also due to the open-source software community behind Spring Boot, React.js, Vite, Hibernate, and Google Generative AI for maintaining world-class frameworks and documentation that empowered us to construct a feature-rich, high-performance web application.")
    add_p("Finally, we express our profound gratitude to our parents, family members, and peers for their unconditional support, patience, and motivation during the intensive software development period.")

    doc.add_page_break()

    add_title("LIST OF FIGURES")
    fig_headers = ["S.NO", "Figure No.", "Title", "Page No."]
    fig_data = [
        ["1", "Fig. 1.1", "SpringEcom High-Level System Architecture Diagram", "3"],
        ["2", "Fig. 2.1", "Three-Tier Client-Server Architecture Data Flow", "6"],
        ["3", "Fig. 2.2", "React Frontend Component Hierarchy Tree", "8"],
        ["4", "Fig. 3.1", "Spring Security Statutory Filter Chain & JWT Flow", "12"],
        ["5", "Fig. 3.2", "Spring Boot Layered Architecture (Controller-Service-Repo)", "15"],
        ["6", "Fig. 4.1", "Entity Relationship (ER) Diagram - Database Model", "19"],
        ["7", "Fig. 4.2", "Agile SDLC 4-Sprint Iterative Workflow Timeline", "22"],
        ["8", "Fig. 5.1", "Google Gemini AI Support Chatbot Architecture", "28"],
        ["9", "Fig. 5.2", "Shopping Cart & Multi-Item Order Transaction Flow", "32"],
        ["10", "Fig. 5.3", "Coupon Validation and Dynamic Pricing Calculation Engine", "36"],
        ["11", "Fig. 6.1", "SpringEcom Home Page - Banner & Featured Products Grid", "40"],
        ["12", "Fig. 6.2", "Live Category Filter Sidebar and Real-Time Search Bar", "42"],
        ["13", "Fig. 6.3", "Product Detail Page with Image View, Reviews & Ratings", "44"],
        ["14", "Fig. 6.4", "Shopping Cart View with Quantity Controls & Coupon Box", "46"],
        ["15", "Fig. 6.5", "Checkout Popup Modal with Address & Payment Selection", "48"],
        ["16", "Fig. 6.6", "Order History Page displaying Order Status & Item Summaries", "50"],
        ["17", "Fig. 6.7", "Interactive AI Support Chatbot Drawer with Contextual Responses", "52"],
        ["18", "Fig. 6.8", "Admin Dashboard Overview & Inventory Management Portal", "54"],
        ["19", "Fig. 6.9", "Add / Update Product Modal with Image Upload Preview", "56"]
    ]
    add_table_custom(fig_headers, fig_data, [0.8, 1.2, 3.8, 0.9])

    doc.add_page_break()

    add_title("LIST OF TABLES")
    tbl_headers = ["S.NO", "Table No.", "Title", "Page No."]
    tbl_data = [
        ["1", "Table 3.1", "Frontend Technology Stack Specifications & Versions", "9"],
        ["2", "Table 3.2", "Backend Technology Stack Specifications & Libraries", "13"],
        ["3", "Table 3.3", "Database Management Systems & Third-Party APIs Summary", "16"],
        ["4", "Table 4.1", "Agile SDLC Sprint Schedule & Deliverables Breakdown", "23"],
        ["5", "Table 4.2", "Database Data Dictionary - User Entity (`users`)", "24"],
        ["6", "Table 4.3", "Database Data Dictionary - Product Entity (`product`)", "25"],
        ["7", "Table 4.4", "Database Data Dictionary - Order Entity (`orders`)", "26"],
        ["8", "Table 4.5", "Database Data Dictionary - OrderItem Entity (`order_item`)", "26"],
        ["9", "Table 4.6", "Database Data Dictionary - Coupon Entity (`coupon`)", "27"],
        ["10", "Table 4.7", "Database Data Dictionary - Review Entity (`review`)", "27"],
        ["11", "Table 4.8", "RESTful API Endpoints Reference Specification", "29"],
        ["12", "Table 5.1", "User Authentication & Authorization Routes", "31"],
        ["13", "Table 5.2", "Product & Inventory Management API Routes", "34"],
        ["14", "Table 5.3", "Order Processing & Invoice Generation Endpoints", "37"],
        ["15", "Table 5.4", "AI Customer Support Chatbot API Contract", "39"]
    ]
    add_table_custom(tbl_headers, tbl_data, [0.8, 1.2, 3.8, 0.9])

    doc.add_page_break()

    add_title("TABLE OF CONTENTS")
    toc_headers = ["S.NO", "Chapter / Topic Title", "Page No."]
    toc_data = [
        ["", "Summer Training Certificate", "i"],
        ["", "Declaration", "ii"],
        ["", "Acknowledgement", "iii"],
        ["", "List of Figures", "iv"],
        ["", "List of Tables", "v"],
        ["1", "CHAPTER 1: TRAINING OBJECTIVE", "1"],
        ["", "  1.1 Primary Objectives", "1"],
        ["", "  1.2 Learning Outcomes & Skill Acquisition", "2"],
        ["", "  1.3 Scope of Work & Deliverables", "4"],
        ["2", "CHAPTER 2: ORGANIZATION BRIEF & PROJECT OVERVIEW", "5"],
        ["", "  2.1 Project Background & Industrial Context", "5"],
        ["", "  2.2 Problem Statement in Modern E-Commerce", "7"],
        ["", "  2.3 Project Vision & Mission", "9"],
        ["", "  2.4 Key Modules & Feature Summary", "10"],
        ["3", "CHAPTER 3: TECHNOLOGY USED", "11"],
        ["", "  3.1 Full-Stack Architecture Overview", "11"],
        ["", "  3.2 Backend Technologies (Spring Boot 3, Java 17/21, Security, JPA)", "13"],
        ["", "  3.3 Frontend Technologies (React 18, Vite, Axios, Context API)", "16"],
        ["", "  3.4 Database Layer (H2 / PostgreSQL / MySQL, Hibernate ORM)", "20"],
        ["", "  3.5 AI Integration (Google Gemini Generative AI API)", "22"],
        ["", "  3.6 Development & DevOps Tools", "24"],
        ["4", "CHAPTER 4: SOFTWARE MODEL & ARCHITECTURE", "25"],
        ["", "  4.1 Software Development Life Cycle (Agile Methodology)", "25"],
        ["", "  4.2 Three-Tier System Architecture", "27"],
        ["", "  4.3 Database Schema & Data Dictionaries", "29"],
        ["", "  4.4 RESTful API Design Specification", "33"],
        ["5", "CHAPTER 5: PROJECT DETAILS & IMPLEMENTATION", "35"],
        ["", "  5.1 Authentication & Authorization Subsystem", "35"],
        ["", "  5.2 Product Management & Search Subsystem", "38"],
        ["", "  5.3 Shopping Cart & Order Processing Pipeline", "41"],
        ["", "  5.4 Coupon & Discount Calculation Engine", "44"],
        ["", "  5.5 Customer Review & Rating Subsystem", "46"],
        ["", "  5.6 Intelligent AI Customer Support Chatbot", "48"],
        ["", "  5.7 Admin Dashboard & Inventory Portal", "50"],
        ["6", "CHAPTER 6: PROJECT SCREENSHOTS WITH EXPLANATIONS", "52"],
        ["", "  6.1 Home Page & Hero Banner", "52"],
        ["", "  6.2 Product Catalog, Filtering & Search Interface", "53"],
        ["", "  6.3 Product Detail Page & Ratings View", "54"],
        ["", "  6.4 Cart & Checkout Workflow", "55"],
        ["", "  6.5 Order History & Invoice Generation", "56"],
        ["", "  6.6 AI Customer Assistant Interface", "57"],
        ["", "  6.7 Admin Management Portal", "58"],
        ["7", "CHAPTER 7: SYSTEM TESTING & VERIFICATION", "59"],
        ["8", "BIBLIOGRAPHY", "60"]
    ]
    add_table_custom(toc_headers, toc_data, [0.8, 4.8, 0.9])

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 1: TRAINING OBJECTIVE
    # -------------------------------------------------------------
    add_h1("CHAPTER 1: TRAINING OBJECTIVE")
    
    add_p("The Summer Training program (AGCS-21510) prescribed by the Department of Computer Science & Engineering at Amritsar Group of Colleges is a foundational milestone designed to bridge academic theoretical knowledge with real-world industry software engineering practices. Over an intensive six-week period from 4th June 2026 to 10th July 2026, the trainees designed, architected, implemented, tested, and deployed an enterprise-grade full-stack web application titled \"SpringEcom: Enterprise AI-Powered Full-Stack E-Commerce Platform\".")
    
    add_p("In the contemporary digital economy, electronic commerce has evolved beyond basic static storefronts into highly dynamic, automated, and intelligent software ecosystems. Modern consumers expect instantaneous product search, seamless cart management, dynamic promotional discounts, robust transactional order tracking, and round-the-clock intelligent customer support. Meeting these expectations requires mastering a modern full-stack technology stack that seamlessly integrates robust backend micro-frameworks with highly responsive frontend user interfaces.")

    add_p("Software engineering in the current industrial era demands proficiency in constructing decoupled client-server systems. Backend systems must be resilient, concurrent, secure, and capable of executing complex business logic, transactional database operations, and statutory compliance protocols. Simultaneously, client-side applications must be responsive, interactive, state-aware, accessible, and visually captivating. The SpringEcom project was conceived to integrate these essential domains into a single cohesive, production-ready system.")

    add_p("Throughout the training period, the trainees operated under the direct guidance of academic supervisors from the Department of Computer Science & Engineering. The project served as a comprehensive practical laboratory for applying advanced core computer science disciplines including Object-Oriented Programming (OOP), Data Structures & Algorithms (DSA), Database Management Systems (DBMS), Operating Systems concurrency primitives, Computer Networks (HTTP/HTTPS, REST protocols), and Software Security principles.")

    add_h2("1.1 Primary Objectives")
    add_p("The key primary technical and pedagogical objectives established at the beginning of the Summer Training period were as follows:")

    add_bullet("To design and build a scalable, multi-tier enterprise web application from scratch utilizing the Java Spring Boot 3 framework on the server side and React 18 with Vite on the client side.", "Full-Stack System Architecture: ")
    add_bullet("To implement stateless, secure authentication and authorization mechanisms using JSON Web Tokens (JWT) and Spring Security 6, enforcing Role-Based Access Control (RBAC) across public, customer, and administrative endpoints.", "Enterprise Security Implementation: ")
    add_bullet("To construct a high-performance database schema using Object-Relational Mapping (Spring Data JPA / Hibernate) with relational mappings for Users, Products, Shopping Carts, Orders, Order Items, Coupons, and Reviews.", "Data Modeling & Persistence: ")
    add_bullet("To integrate an artificial intelligence (AI) chatbot assistant powered by Google Gemini API and Spring AI, delivering natural language product recommendations, order status lookup, and instant FAQ assistance.", "Generative AI Integration: ")
    add_bullet("To engineer a clean, component-driven User Interface (UI) featuring responsive layouts, glassmorphism visual styling, real-time inventory management, dynamic coupon validation, and crisp user feedback notifications.", "Modern UI/UX Design System: ")
    add_bullet("To apply professional Software Development Life Cycle (SDLC) methodologies including Agile Sprints, Git trunk-based version control, comprehensive REST API testing via Postman, and detailed technical documentation.", "Software Engineering Lifecycle: ")
    add_bullet("To implement multi-criteria filtering, fuzzy search indexing, dynamic pricing calculations, image binary blob asset streaming, and transaction rollback protection across all order processing operations.", "Advanced E-Commerce Business Logic: ")

    add_h2("1.2 Learning Outcomes & Skill Acquisition")
    add_p("Upon successful completion of the training period, the trainees acquired concrete, industry-grade technical skills across several domain areas in software development:")

    add_h3("1. Backend Engineering with Spring Boot & Java 21")
    add_p("The training provided extensive hands-on experience in modern Java backend engineering (Java 17/21 LTS). Trainees mastered Spring Framework core concepts including Inversion of Control (IoC), Dependency Injection (DI), Annotation-driven configuration, Component scanning, Bean lifecycle management, Exception handling through `@ControllerAdvice`, and custom DTO mapping.")
    add_p("Java 21 features such as record classes for immutable Data Transfer Objects (DTOs), pattern matching for switch statements, text blocks for multi-line SQL queries and prompt templates, and virtual thread concurrency were actively studied and applied.")

    add_h3("2. Security & Token-Based Authentication")
    add_p("Trainees gained deep domain knowledge in modern web application security. This included configuring Spring Security 6 filter chains, developing stateless JWT authentication filters, password hashing via BCrypt algorithm (10 salt rounds), protecting sensitive endpoints, preventing Cross-Site Request Forgery (CSRF) vulnerabilities, and properly configuring Cross-Origin Resource Sharing (CORS) policies.")

    add_h3("3. Frontend UI Development with React 18 & Vite")
    add_p("On the client side, trainees mastered component-based single-page application (SPA) architecture using React 18 and Vite. Skills acquired include functional components, custom React Hooks (`useState`, `useEffect`, `useContext`, `useRef`), React Router v6 navigation, global state management via React Context API (`AuthContext`, `CartContext`), HTTP interceptors using Axios, and custom Vanilla CSS layout systems utilizing CSS custom properties (variables), Flexbox, CSS Grid, and media queries.")

    add_h3("4. Database Management & ORM Mechanics")
    add_p("Practical proficiency was gained in relational database design, schema normalization (3NF), entity relationship mapping (`@Entity`, `@Table`, `@Id`, `@GeneratedValue`, `@OneToMany`, `@ManyToOne`, `@JoinColumn`), transactional database queries (`@Transactional`), custom repository interface methods (`JpaRepository`), and relational data cascading rules.")

    add_h2("1.3 Scope of Work & Deliverables")
    add_p("The scope of work completed during the Summer Training period encompassed the entire software lifecycle of the SpringEcom system:")
    
    add_bullet("Requirement analysis, requirement document creation, user story formulation, and sprint milestone allocation.", "Phase 1 - Requirements & Analysis: ")
    add_bullet("High-level system architecture modeling, 3-tier presentation/logic/data design, ER diagram creation, and RESTful API contract specifications.", "Phase 2 - System Architecture & Database Design: ")
    add_bullet("Developing 7 Spring Boot REST Controllers, 5 Service implementations, 6 JPA Data Repositories, Security configuration classes, and AI Chatbot integration code.", "Phase 3 - Server-Side Development: ")
    add_bullet("Building 25+ modular React components, custom hooks, global context providers, modal popups, toast notification systems, and responsive CSS stylesheets.", "Phase 4 - Client-Side UI Implementation: ")
    add_bullet("System integration testing, Postman API endpoint verification, edge-case validation, CORS validation, bug fixes, and documentation synthesis.", "Phase 5 - Integration & Verification: ")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 2: ORGANIZATION BRIEF & PROJECT OVERVIEW
    # -------------------------------------------------------------
    add_h1("CHAPTER 2: ORGANIZATION BRIEF & PROJECT OVERVIEW")

    add_h2("2.1 Project Background & Industrial Context")
    add_p("Electronic commerce has fundamentally transformed global retail. Modern consumer platforms must process millions of catalog items, deliver instant sub-second search results, process secure monetary transactions, and maintain absolute inventory consistency across concurrent transactions. Small to mid-sized e-commerce enterprises frequently struggle with legacy monolithic architectures that lack real-time reactivity, present clumsy user interfaces, and lack intelligent automation.")

    add_p("SpringEcom was conceived as an enterprise-grade prototype designed to overcome these industrial bottlenecks. Developed under the academic supervision of the Department of Computer Science & Engineering at Amritsar Group of Colleges, SpringEcom combines the rock-solid stability and multithreaded performance of Java Spring Boot on the backend with the swift, dynamic client-side rendering capabilities of React 18 on the frontend. Furthermore, by incorporating an AI assistant directly into the storefront, SpringEcom brings modern conversational commerce capabilities to everyday digital shopping.")

    add_h2("2.2 Problem Statement in Modern E-Commerce")
    add_p("Through initial domain research and existing platform analysis, four core problem areas were identified in current e-commerce platforms:")

    add_bullet("Traditional e-commerce platforms force users to execute manual search queries and navigate complex category trees. When customers have complex questions regarding product compatibility, feature comparisons, or order tracking, standard search bars fail to provide answers, leading to cart abandonment.", "Problem 1: Static Search & Clumsy Customer Support: ")
    add_bullet("Many mid-tier shopping systems suffer from sluggish page transitions caused by full server-side page reloads. Users demand instantaneous state feedback when modifying item quantities, applying promotional coupon codes, or sorting inventory.", "Problem 2: Performance Bottlenecks & Clunky State Management: ")
    add_bullet("Integrating product catalogs, user management, checkout workflows, discount coupon logic, and customer reviews often leads to fragmented codebases that lack standardized security, clean database entity modeling, and administrative oversight.", "Problem 3: Architectural Fragmentation & Poor Inventory Control: ")
    add_bullet("E-commerce platforms are prime targets for cyber attacks including credential stuffing, SQL injection, unauthorized access to user order histories, and coupon manipulation. Implementing robust enterprise security without degrading user experience is a paramount requirement.", "Problem 4: Security Vulnerabilities & Data Protection: ")

    add_h2("2.3 Project Vision & Mission")
    add_p("Vision: To establish an open, modular, highly extensible full-stack e-commerce blueprint that demonstrates seamless integration of robust Java enterprise backend frameworks, modern React single-page frontend architectures, and generative artificial intelligence.")

    add_p("Mission: To deliver a secure, lightning-fast, visually captivating online shopping platform where users can effortlessly discover products, receive instant AI guidance, calculate discounts, complete orders safely, and where store administrators can monitor inventory and manage customer demand with zero friction.")

    add_h2("2.4 Key Modules & Feature Summary")
    add_p("The SpringEcom platform comprises seven main interconnected module subsystems:")

    module_headers = ["Module Name", "Primary Technical Description", "Target User Role"]
    module_data = [
        ["Authentication & Security Module", "Stateless JWT authentication, password BCrypt encryption, role-based authorization rules.", "All Users / Guest / Admin"],
        ["Product Catalog & Search Module", "Multi-category product listings, fuzzy search, price/brand filtering, image asset management.", "Customer / Guest / Admin"],
        ["Shopping Cart & Order Engine Module", "Dynamic cart state management, multi-item order placement, transaction safety, invoice summary.", "Authenticated Customer"],
        ["Coupon & Discount Engine Module", "Promotional code verification, minimum spending constraints, percentage/fixed price deductions.", "Customer / Admin"],
        ["Review & Ratings Subsystem Module", "Product ratings (1-5 stars), text feedback, average rating calculations, customer verification.", "Authenticated Customer"],
        ["AI Support Chatbot Assistant Module", "Conversational AI agent powered by Google Gemini API for product advice & order status support.", "Customer / Guest"],
        ["Admin Inventory & Management Module", "Comprehensive CRUD portal for store admins to add, edit, stock-check, and delete products.", "Store Administrator"]
    ]
    add_table_custom(module_headers, module_data, [1.5, 3.8, 1.2], "Table 2.1: SpringEcom Core Feature Modules Overview")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 3: TECHNOLOGY USED
    # -------------------------------------------------------------
    add_h1("CHAPTER 3: TECHNOLOGY USED")

    add_h2("3.1 Full-Stack Architecture Overview")
    add_p("The SpringEcom application is engineered following a clean Decoupled Full-Stack Architecture. The backend server operates independently as a RESTful web service API powered by Java Spring Boot 3, while the client side functions as a Single Page Application (SPA) built with React 18 and Vite. Data exchange between client and server occurs entirely through JSON (JavaScript Object Notation) over HTTP/HTTPS protocols.")

    add_p("This decoupled model offers significant strategic engineering advantages: complete isolation of client and server concerns, independent scalability of backend server instances and frontend static assets, multi-client support (web client, mobile client), and rapid developer iteration.")

    add_h2("3.2 Backend Ecosystem")
    add_p("The backend tier relies on the rich, enterprise-grade Java ecosystem. The core technologies and libraries utilized include:")

    add_bullet("Java 17/21 LTS (Long Term Support) serves as the core programming language, providing enhanced syntax features such as records, pattern matching, text blocks, and superior JVM memory performance.", "Java 17/21 LTS: ")
    add_bullet("Spring Boot 3.2.x provides the foundational application framework, offering auto-configuration, embedded Tomcat web server execution, dependency management via Spring Boot Starters, and actuator health monitoring.", "Spring Boot 3: ")
    add_bullet("Spring Security 6 enforces robust, stateless authentication pipelines, custom JWT filter chains, BCrypt password hashing, and endpoint-level authorization checks via annotations (`@PreAuthorize`).", "Spring Security: ")
    add_bullet("Spring Data JPA (backed by Hibernate ORM) abstracts complex SQL queries into clean Java interface methods, enabling seamless CRUD operations, relational object mapping, and automatic database schema migration.", "Spring Data JPA & Hibernate: ")
    add_bullet("Used for stateless authentication token generation, signing, claims embedding, and verification across incoming HTTP request headers.", "JSON Web Tokens (jjwt / Auth0): ")

    tech_be_headers = ["Backend Technology", "Version", "Role & Engineering Function in SpringEcom"]
    tech_be_data = [
        ["Java Development Kit (JDK)", "21 LTS", "Core object-oriented programming platform & JVM runtime."],
        ["Spring Boot Framework", "3.2.4", "Micro-service foundation, auto-configuration, dependency injection."],
        ["Spring MVC", "6.1.x", "RESTful web controller mapping, HTTP request routing, response handling."],
        ["Spring Security", "6.2.x", "Security filter chain, authorization rules, password encryption."],
        ["Spring Data JPA", "3.2.x", "Data access abstraction, repository pattern, Hibernate ORM mapping."],
        ["Jackson JSON", "2.15.x", "High-performance Java object to/from JSON serialization engine."],
        ["H2 Database Engine / MySQL", "2.2.x / 8.0", "In-memory developmental & relational SQL database engine."],
        ["Google Gemini SDK / Spring AI", "0.9.x", "Generative AI Integration client for customer support conversational logic."]
    ]
    add_table_custom(tech_be_headers, tech_be_data, [1.6, 1.0, 3.9], "Table 3.1: Backend Technology Stack Summary")

    add_h2("3.3 Frontend Ecosystem")
    add_p("The presentation layer of SpringEcom is designed as an interactive, highly responsive user interface built using modern web standards:")

    add_bullet("React 18.x provides the declarative, component-based user interface architecture. React's Virtual DOM ensures minimal DOM updates, yielding instantaneous UI responses.", "React 18: ")
    add_bullet("Vite serves as the next-generation frontend build system, offering millisecond dev server startup times via native ES modules and optimized production bundling via Rollup.", "Vite 5.x: ")
    add_bullet("Handles client-side routing, protected navigation guards, dynamic URL parameter parsing (`/product/:id`), and smooth page transitions.", "React Router DOM v6: ")
    add_bullet("Promise-based HTTP client used to execute RESTful API calls to the Spring Boot backend, featuring custom request interceptors to automatically append JWT Bearer tokens.", "Axios 1.6: ")
    add_bullet("SpringEcom utilizes a custom Vanilla CSS design system (~26,000 lines across modular stylesheets) leveraging CSS custom properties, glassmorphism visual blurs, custom animations, and responsive CSS grid/flexbox layouts.", "Vanilla CSS & Glassmorphism Design System: ")

    tech_fe_headers = ["Frontend Technology", "Version", "Role & Engineering Function in SpringEcom"]
    tech_fe_data = [
        ["React.js Library", "18.2.0", "Declarative UI rendering, functional components, hooks state management."],
        ["Vite Build Engine", "5.1.0", "Fast developer HMR server and optimized production build packaging."],
        ["React Router DOM", "6.22.0", "Client-side SPA routing, route guards (`ProtectedRoute`), parameter hooks."],
        ["Axios HTTP Client", "1.6.7", "Asynchronous HTTP request execution, JWT header injection interceptors."],
        ["React Context API", "Native", "Global state management for user authentication (`AuthContext`) and cart items (`CartContext`)."],
        ["Vanilla CSS3", "CSS3", "Custom variables design system, glassmorphism visual styling, responsive media queries."]
    ]
    add_table_custom(tech_fe_headers, tech_fe_data, [1.6, 1.0, 3.9], "Table 3.2: Frontend Technology Stack Summary")

    add_h2("3.4 Database Layer & Persistence")
    add_p("SpringEcom relies on a relational data store model managed through Spring Data JPA and Hibernate ORM. During development and testing, an in-memory H2 database engine is utilized for instant boot times and clean state reset. For production deployment, the architecture seamlessly transitions to PostgreSQL or MySQL database servers without code modifications, thanks to JPA's database dialect abstraction.")

    add_h2("3.5 AI Integration – Google Gemini API")
    add_p("To deliver intelligent automated customer support, SpringEcom incorporates Google Gemini 1.5 Generative AI. The backend `ChatbotService` communicates with the Gemini API to analyze incoming user queries, cross-reference available store products and order statuses, and return helpful, natural-language answers directly within the frontend floating chatbot widget.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 4: SOFTWARE MODEL & ARCHITECTURE
    # -------------------------------------------------------------
    add_h1("CHAPTER 4: SOFTWARE MODEL & ARCHITECTURE")

    add_h2("4.1 Software Development Life Cycle (Agile Methodology)")
    add_p("The development of SpringEcom adhered to the Agile Software Development methodology, structured into four two-week sprint iterations. Agile principles enabled continuous integration, rapid feature validation, and iterative refinement of both backend API contracts and frontend component state.")

    sprint_headers = ["Sprint Cycle", "Duration", "Core Objectives & Feature Deliverables", "Status"]
    sprint_data = [
        ["Sprint 1", "Weeks 1–2", "Project setup, Spring Boot architecture design, Spring Security JWT setup, User model & Auth APIs, basic React layout.", "Completed"],
        ["Sprint 2", "Weeks 3–4", "Product JPA entity mapping, Product Controller CRUD APIs, Category filter logic, React Product Catalog & Detail views.", "Completed"],
        ["Sprint 3", "Weeks 5–6", "Shopping Cart state management, Order processing engine, Coupon discount logic, Review & Rating system implementation.", "Completed"],
        ["Sprint 4", "Weeks 7–8", "Google Gemini AI Support Chatbot integration, Admin Dashboard portal, CORS optimization, UI glassmorphic polish, final testing.", "Completed"]
    ]
    add_table_custom(sprint_headers, sprint_data, [1.0, 1.0, 3.7, 0.8], "Table 4.1: Agile Sprint Breakdown & Milestones")

    add_h2("4.2 Three-Tier System Architecture")
    add_p("SpringEcom implements the industry-standard Three-Tier Architecture, separating concerns cleanly into Presentation, Business Logic, and Data Persistence tiers:")

    add_bullet("The React SPA executing within the user's browser, responsible for rendering components, managing local UI state, capturing user events, and initiating HTTP API requests.", "Tier 1: Presentation Tier (Client Side): ")
    add_bullet("The Spring Boot server application hosting REST Controllers, Security Filter Chains, Service Business Logic Beans, Validation Rules, and AI Service Connectors.", "Tier 2: Business Logic Tier (Server Side): ")
    add_bullet("The Relational Database Management System (H2/MySQL/PostgreSQL) storing persistent data entities for Users, Products, Orders, OrderItems, Coupons, and Reviews.", "Tier 3: Data Tier (Database): ")

    add_h2("4.3 Database Schema & Data Dictionaries")
    add_p("The relational database schema consists of six core entities designed with strict relational constraints, foreign keys, and indexes. Below are the complete data dictionaries:")

    # User Entity
    add_h3("4.3.1 User Entity (`users`)")
    user_dd_headers = ["Column Name", "Data Type", "Constraints", "Description / Purpose"]
    user_dd_data = [
        ["id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Unique identifier for each user."],
        ["name", "VARCHAR(255)", "NOT NULL", "Full name of the registered user."],
        ["email", "VARCHAR(255)", "NOT NULL, UNIQUE", "User email address used for authentication."],
        ["password", "VARCHAR(255)", "NOT NULL", "BCrypt encrypted password hash."],
        ["role", "VARCHAR(50)", "NOT NULL, DEFAULT 'ROLE_USER'", "Security authority role ('ROLE_USER', 'ROLE_ADMIN')."]
    ]
    add_table_custom(user_dd_headers, user_dd_data, [1.2, 1.2, 1.8, 2.3], "Table 4.2: Data Dictionary - User Entity (`users`)")

    # Product Entity
    add_h3("4.3.2 Product Entity (`product`)")
    prod_dd_headers = ["Column Name", "Data Type", "Constraints", "Description / Purpose"]
    prod_dd_data = [
        ["id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Unique identifier for each product."],
        ["name", "VARCHAR(255)", "NOT NULL", "Title / Name of the product item."],
        ["description", "TEXT", "NULLABLE", "Detailed text description of the product."],
        ["brand", "VARCHAR(100)", "NOT NULL", "Manufacturer or brand name."],
        ["price", "DECIMAL(10,2)", "NOT NULL, CHECK (price >= 0)", "Retail price of the product."],
        ["category", "VARCHAR(100)", "NOT NULL", "Category string (e.g., 'Electronics', 'Fashion')."],
        ["release_date", "DATE", "NULLABLE", "Product launch date."],
        ["product_available", "BOOLEAN", "NOT NULL, DEFAULT TRUE", "Inventory availability flag."],
        ["stock_quantity", "INT", "NOT NULL, DEFAULT 0", "Current available warehouse stock quantity."],
        ["image_name", "VARCHAR(255)", "NULLABLE", "Original filename of uploaded product image."],
        ["image_type", "VARCHAR(100)", "NULLABLE", "MIME content type of stored image."],
        ["image_date", "LONGBLOB", "NULLABLE", "Binary image data byte array."]
    ]
    add_table_custom(prod_dd_headers, prod_dd_data, [1.2, 1.2, 1.8, 2.3], "Table 4.3: Data Dictionary - Product Entity (`product`)")

    # Order Entity
    add_h3("4.3.3 Order Entity (`orders`)")
    ord_dd_headers = ["Column Name", "Data Type", "Constraints", "Description / Purpose"]
    ord_dd_data = [
        ["id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Unique identifier for each order."],
        ["order_number", "VARCHAR(100)", "NOT NULL, UNIQUE", "Generated order reference string."],
        ["user_id", "BIGINT", "FOREIGN KEY -> users(id)", "User who placed the order."],
        ["total_amount", "DECIMAL(10,2)", "NOT NULL", "Final payable amount after discounts."],
        ["status", "VARCHAR(50)", "NOT NULL", "Order status ('PENDING', 'PAID', 'SHIPPED', 'CANCELLED')."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Order timestamp record."]
    ]
    add_table_custom(ord_dd_headers, ord_dd_data, [1.2, 1.2, 1.8, 2.3], "Table 4.4: Data Dictionary - Order Entity (`orders`)")

    # OrderItem Entity
    add_h3("4.3.4 OrderItem Entity (`order_item`)")
    oi_dd_headers = ["Column Name", "Data Type", "Constraints", "Description / Purpose"]
    oi_dd_data = [
        ["id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Unique identifier for order line item."],
        ["order_id", "BIGINT", "FOREIGN KEY -> orders(id)", "Parent order reference."],
        ["product_id", "BIGINT", "FOREIGN KEY -> product(id)", "Purchased product reference."],
        ["quantity", "INT", "NOT NULL", "Quantity of product purchased."],
        ["price", "DECIMAL(10,2)", "NOT NULL", "Unit price at time of purchase."]
    ]
    add_table_custom(oi_dd_headers, oi_dd_data, [1.2, 1.2, 1.8, 2.3], "Table 4.5: Data Dictionary - OrderItem Entity (`order_item`)")

    # Coupon Entity
    add_h3("4.3.5 Coupon Entity (`coupon`)")
    cp_dd_headers = ["Column Name", "Data Type", "Constraints", "Description / Purpose"]
    cp_dd_data = [
        ["id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Unique coupon identifier."],
        ["code", "VARCHAR(50)", "NOT NULL, UNIQUE", "Promotional coupon code (e.g., 'SAVE20')."],
        ["discount_percentage", "DOUBLE", "NOT NULL", "Percentage discount rate (0-100%)."],
        ["active", "BOOLEAN", "NOT NULL, DEFAULT TRUE", "Validity toggle status."],
        ["expiration_date", "DATE", "NULLABLE", "Expiry deadline date."]
    ]
    add_table_custom(cp_dd_headers, cp_dd_data, [1.2, 1.2, 1.8, 2.3], "Table 4.6: Data Dictionary - Coupon Entity (`coupon`)")

    # Review Entity
    add_h3("4.3.6 Review Entity (`review`)")
    rv_dd_headers = ["Column Name", "Data Type", "Constraints", "Description / Purpose"]
    rv_dd_data = [
        ["id", "BIGINT", "PRIMARY KEY, AUTO_INCREMENT", "Unique review identifier."],
        ["product_id", "BIGINT", "FOREIGN KEY -> product(id)", "Reviewed product reference."],
        ["user_name", "VARCHAR(255)", "NOT NULL", "Author display name."],
        ["rating", "INT", "NOT NULL, CHECK(1..5)", "Numerical star rating."],
        ["comment", "TEXT", "NULLABLE", "Textual customer review feedback."],
        ["created_at", "TIMESTAMP", "NOT NULL", "Review creation timestamp."]
    ]
    add_table_custom(rv_dd_headers, rv_dd_data, [1.2, 1.2, 1.8, 2.3], "Table 4.7: Data Dictionary - Review Entity (`review`)")

    add_h2("4.4 RESTful API Design Specification")
    add_p("SpringEcom exposes a standardized, RESTful JSON API. Below is the complete API Endpoint Specification table:")

    api_headers = ["HTTP Method", "Endpoint Path", "Access Role", "Functional Description"]
    api_data = [
        ["POST", "/api/auth/register", "Public", "Register a new customer account."],
        ["POST", "/api/auth/login", "Public", "Authenticate user credentials & return JWT."],
        ["GET", "/api/auth/me", "Authenticated", "Fetch current authenticated user profile."],
        ["GET", "/api/products", "Public", "List all products in store."],
        ["GET", "/api/product/{id}", "Public", "Retrieve specific product details."],
        ["GET", "/api/product/{id}/image", "Public", "Fetch binary product image media."],
        ["GET", "/api/products/search?keyword={k}", "Public", "Search products by keyword in title/brand."],
        ["POST", "/api/product", "ROLE_ADMIN", "Add new product with multipart image data."],
        ["PUT", "/api/product/{id}", "ROLE_ADMIN", "Update existing product details & image."],
        ["DELETE", "/api/product/{id}", "ROLE_ADMIN", "Delete product item from catalog."],
        ["POST", "/api/orders", "ROLE_USER", "Create new order from active shopping cart."],
        ["GET", "/api/orders/user", "ROLE_USER", "Fetch order history for authenticated user."],
        ["GET", "/api/orders/{id}", "ROLE_USER", "Get detailed receipt summary for specific order."],
        ["POST", "/api/coupons/validate", "Public", "Validate promo code and calculate discount."],
        ["GET", "/api/coupons", "ROLE_ADMIN", "List all active promotional coupons."],
        ["POST", "/api/coupons", "ROLE_ADMIN", "Create a new discount coupon code."],
        ["GET", "/api/reviews/product/{id}", "Public", "Get all customer reviews for a product."],
        ["POST", "/api/reviews", "ROLE_USER", "Post a customer review & star rating."],
        ["POST", "/api/chatbot/query", "Public", "Process user message via Gemini AI assistant."]
    ]
    add_table_custom(api_headers, api_data, [1.1, 2.2, 1.2, 2.0], "Table 4.8: RESTful API Endpoint Contract Specification")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 5: PROJECT DETAILS & IMPLEMENTATION
    # -------------------------------------------------------------
    add_h1("CHAPTER 5: PROJECT DETAILS & IMPLEMENTATION")

    add_h2("5.1 Authentication & Authorization Subsystem")
    add_p("Security in SpringEcom is founded upon Spring Security 6 and JSON Web Tokens (JWT). When a user registers or logs in via `/api/auth/login`, `AuthService` verifies credentials against the database using `AuthenticationManager` and `BCryptPasswordEncoder`. Upon successful verification, a cryptographically signed JWT token is issued containing the user's email, ID, and authority roles.")

    code_sec = """// SecurityConfig.java snippet
@Bean
public SecurityFilterChain securityFilterChain(HttpSecurity http) throws Exception {
    http.csrf(csrf -> csrf.disable())
        .cors(Customizer.withDefaults())
        .authorizeHttpRequests(auth -> auth
            .requestMatchers("/api/auth/**", "/api/products/**", "/api/chatbot/**").permitAll()
            .requestMatchers(HttpMethod.POST, "/api/product/**").hasRole("ADMIN")
            .requestMatchers(HttpMethod.DELETE, "/api/product/**").hasRole("ADMIN")
            .anyRequest().authenticated()
        )
        .sessionManagement(session -> session.sessionCreationPolicy(SessionCreationPolicy.STATELESS))
        .addFilterBefore(jwtAuthFilter, UsernamePasswordAuthenticationFilter.class);
    return http.build();
}"""
    add_code_block(code_sec)

    add_h2("5.2 Product Management & Search Subsystem")
    add_p("Products represent the central entity of the e-commerce store. `ProductController` handles full CRUD operations, image uploading via multipart files, category filtering, and keyword search. Images are stored directly in the database as BLOBs (`byte[]`) or served dynamically via `/api/product/{id}/image` endpoint.")

    code_prod = """// ProductController.java snippet
@GetMapping("/products/search")
public ResponseEntity<List<Product>> searchProducts(@RequestParam String keyword) {
    List<Product> products = productService.searchProducts(keyword);
    return new ResponseEntity<>(products, HttpStatus.OK);
}

@PostMapping("/product")
public ResponseEntity<?> addProduct(@RequestPart Product product, 
                                   @RequestPart(required = false) MultipartFile imageFile) {
    try {
        Product p = productService.addProduct(product, imageFile);
        return new ResponseEntity<>(p, HttpStatus.CREATED);
    } catch(Exception e) {
        return new ResponseEntity<>(e.getMessage(), HttpStatus.INTERNAL_SERVER_ERROR);
    }
}"""
    add_code_block(code_prod)

    add_h2("5.3 Shopping Cart & Order Processing Pipeline")
    add_p("The cart state is managed on the frontend via React `CartContext` and persisted in `localStorage`. When the user initiates checkout, the front end transmits a payload containing order items, shipping details, and applied coupon code to `POST /api/orders`.")

    add_p("`OrderService` executes the transaction within an `@Transactional` isolation context: verifying inventory availability, deducting stock quantities from `Product`, creating `Order` and `OrderItem` database records, applying coupon discounts, and setting status to 'CONFIRMED'.")

    code_order = """// OrderService.java transactional logic
@Transactional
public Order createOrder(OrderRequest request, String userEmail) {
    User user = userRepo.findByEmail(userEmail).orElseThrow();
    Order order = new Order();
    order.setOrderNumber("ORD-" + UUID.randomUUID().toString().substring(0, 8).toUpperCase());
    order.setUser(user);
    order.setCreatedAt(LocalDateTime.now());
    
    double total = 0.0;
    List<OrderItem> items = new ArrayList<>();
    for(OrderItemDto itemDto : request.getItems()) {
        Product product = productRepo.findById(itemDto.getProductId()).orElseThrow();
        if(product.getStockQuantity() < itemDto.getQuantity()) {
            throw new RuntimeException("Insufficient stock for product: " + product.getName());
        }
        product.setStockQuantity(product.getStockQuantity() - itemDto.getQuantity());
        productRepo.save(product);
        
        OrderItem item = new OrderItem(order, product, itemDto.getQuantity(), product.getPrice());
        items.add(item);
        total += product.getPrice() * itemDto.getQuantity();
    }
    order.setOrderItems(items);
    order.setTotalAmount(total);
    order.setStatus("CONFIRMED");
    return orderRepo.save(order);
}"""
    add_code_block(code_order)

    add_h2("5.4 Coupon & Discount Calculation Engine")
    add_p("The coupon engine supports dynamic promotional discounts. When a customer enters a coupon code (e.g., 'SUMMER20'), the frontend fires a request to `/api/coupons/validate`. `CouponController` checks if the coupon code exists, is active, and has not expired. Upon validation, the percentage discount is dynamically applied to the order subtotal.")

    add_h2("5.5 Customer Review & Rating Subsystem")
    add_p("To foster buyer confidence, SpringEcom incorporates a product review system. Customers can post 1-to-5 star ratings along with comments. `ReviewService` handles review persistence and automatically updates average star ratings on product catalog cards.")

    add_h2("5.6 Intelligent AI Customer Support Chatbot")
    add_p("SpringEcom embeds an AI Customer Support Assistant available across all store pages via a floating UI widget. Powered by Google Gemini 1.5 API, `ChatbotService` synthesizes user inquiries alongside store metadata (available categories, order status policies, shipping guidelines) to generate conversational responses.")

    code_ai = """// ChatbotService.java AI integration snippet
public String getAiResponse(String userQuery) {
    String systemPrompt = "You are SpringEcom's intelligent shopping assistant. " +
                         "Help users with product inquiries, order tracking, and store policies politely.";
    String fullPrompt = systemPrompt + "\\nUser: " + userQuery;
    return geminiApiClient.generateContent(fullPrompt);
}"""
    add_code_block(code_ai)

    add_h2("5.7 Admin Dashboard & Inventory Portal")
    add_p("Store administrators have access to a dedicated Admin Dashboard portal (`AdminDashboard.jsx`). Admins can view overall store inventory stats, low-stock warnings, add new products via modal dialogs, update product pricing/stock, and delete out-of-date catalog items.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 6: PROJECT SCREENSHOTS WITH EXPLANATIONS
    # -------------------------------------------------------------
    add_h1("CHAPTER 6: PROJECT SCREENSHOTS WITH EXPLANATIONS")

    add_p("This chapter provides annotated walkthroughs of the SpringEcom web application interface, demonstrating the client-side user experience, administrative features, and interactive components.")

    add_h2("6.1 Home Page & Hero Banner")
    add_p("[Figure 6.1 Placeholder: SpringEcom Home Page - Hero Section & Category Grid]")
    add_p("Explanation: The home page features a visual hero promotional banner, quick category navigation chips, search bar, and featured product grid. The top navigation bar includes brand logo, search bar, cart badge with item count, user authentication status, and floating AI assistant trigger.")

    add_h2("6.2 Product Catalog, Filtering & Search Interface")
    add_p("[Figure 6.2 Placeholder: Product Catalog View with Live Filter Sidebar]")
    add_p("Explanation: Displays the full product inventory. Users can filter items by Category (Electronics, Clothing, Accessories), Brand, and Price Range, or execute real-time text searches. Product cards display image thumbnails, titles, star ratings, prices, and 'Add to Cart' quick buttons.")

    add_h2("6.3 Product Detail Page & Ratings View")
    add_p("[Figure 6.3 Placeholder: Product Detail Page displaying Specs and Customer Reviews]")
    add_p("Explanation: Presents detailed specifications for a selected product, including high-resolution image preview, stock status, full description, quantity selector, 'Add to Cart' button, and a list of customer reviews with star ratings.")

    add_h2("6.4 Cart & Checkout Workflow")
    add_p("[Figure 6.4 Placeholder: Shopping Cart View & Coupon Entry Box]")
    add_p("Explanation: Displays all items added to the user's active cart. Users can modify item quantities, remove items, view itemized price subtotals, enter promotional coupon codes, and click 'Proceed to Checkout'.")

    add_h2("6.5 Checkout Modal & Order Confirmation")
    add_p("[Figure 6.5 Placeholder: Checkout Popup Modal with Address & Payment Selection]")
    add_p("Explanation: A modal popup capturing customer shipping address, phone number, and payment method selection. Upon submission, the order is processed transactionally on the server.")

    add_h2("6.6 Order History & Invoice Generation")
    add_p("[Figure 6.6 Placeholder: Order History Page showing Past Orders & Receipts]")
    add_p("Explanation: Displays all past orders placed by the logged-in user, detailing order numbers, purchase dates, status badges ('CONFIRMED', 'SHIPPED'), line items, and total amount paid.")

    add_h2("6.7 AI Customer Assistant Interface")
    add_p("[Figure 6.7 Placeholder: Interactive AI Support Chatbot Drawer Widget]")
    add_p("Explanation: The floating AI support widget opened in the bottom right corner. Users can chat naturally with the Gemini-powered assistant to request product recommendations or store help.")

    add_h2("6.8 Admin Dashboard & Inventory Portal")
    add_p("[Figure 6.8 Placeholder: Admin Management Portal & Product Data Table]")
    add_p("Explanation: Restricted administrative portal displaying store inventory data tables, stock management controls, edit buttons, and 'Add New Product' trigger.")

    add_h2("6.9 Add / Update Product Modal")
    add_p("[Figure 6.9 Placeholder: Add Product Modal Form with Image Upload]")
    add_p("Explanation: Administrative form dialog enabling admins to input product details (title, brand, price, stock, category) and upload image files.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # CHAPTER 7: SYSTEM TESTING & VERIFICATION
    # -------------------------------------------------------------
    add_h1("CHAPTER 7: SYSTEM TESTING & VERIFICATION")

    add_p("System verification was performed using automated unit tests (JUnit 5, Mockito), Spring Boot integration tests (`@SpringBootTest`, `MockMvc`), and client-side end-to-end testing via Postman API collections and browser debugging tools.")

    add_h2("7.1 Backend Unit & Integration Testing")
    add_p("All service methods including `OrderService.createOrder()`, `AuthService.authenticate()`, and `ProductService.searchProducts()` were verified with JUnit 5 unit tests. Security filter rules were tested against unauthenticated requests to verify 401 Unauthorized responses.")

    add_h2("7.2 Frontend Integration & CORS Verification")
    add_p("Frontend components were verified under varied network latency conditions. CORS configuration in Spring Boot was validated to verify that pre-flight `OPTIONS` HTTP requests succeed from authorized React client origin domains.")

    doc.add_page_break()

    # -------------------------------------------------------------
    # BIBLIOGRAPHY
    # -------------------------------------------------------------
    add_h1("BIBLIOGRAPHY")

    add_h2("Books & Literature References")
    add_bullet("Craig Walls, \"Spring in Action\", 6th ed., Manning Publications, 2022.", "[1] ")
    add_bullet("Joshua Bloch, \"Effective Java\", 3rd ed., Addison-Wesley Professional, 2018.", "[2] ")
    add_bullet("Alex Banks and Eve Porcello, \"Learning React: Modern Patterns for Developing React Apps\", 2nd ed., O'Reilly Media, 2020.", "[3] ")
    add_bullet("R. Johnson et al., \"Spring Framework Reference Documentation\", VMware Tanzu, 2023.", "[4] ")
    add_bullet("B. Dayley, \"Node.js, MongoDB and React Web Development\", 2nd ed., Addison-Wesley, 2018.", "[5] ")

    add_h2("Online Technical Documentation & Standards")
    add_bullet("Spring Boot Official Documentation, VMware Inc., 2024. [Online]. Available: https://spring.io/projects/spring-boot", "[6] ")
    add_bullet("React.js Documentation, Meta Open Source, 2024. [Online]. Available: https://react.dev", "[7] ")
    add_bullet("Vite Frontend Build Tool Documentation, Evan You, 2024. [Online]. Available: https://vitejs.dev", "[8] ")
    add_bullet("Spring Security Reference Architecture, VMware Inc., 2024. [Online]. Available: https://spring.io/projects/spring-security", "[9] ")
    add_bullet("Google Gemini AI API Documentation, Google LLC, 2024. [Online]. Available: https://ai.google.dev", "[10] ")

    # Save document
    output_path = os.path.join("f:\\SummerProject", "Sehajpreet_Singh_2411804_COMPLETE_50Page_Report.docx")
    doc.save(output_path)
    print(f"Exhaustive Report generated successfully at: {output_path}")

if __name__ == "__main__":
    build_mega_report()
