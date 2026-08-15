#!/usr/bin/env python3
"""
Script to edit the Spring Boot E-Commerce Report
- Replace all student names with Sehajpreet Singh (2411804)
- Remove all code sections
- Add screenshot placeholders
- Keep everything else intact
"""

import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def edit_report(input_file, output_file):
    """Edit the report with correct details"""
    
    print(f"Opening: {input_file}")
    doc = docx.Document(input_file)
    
    changes_made = 0
    code_sections_removed = 0
    
    # Process all paragraphs
    for para in doc.paragraphs:
        original_text = para.text
        
        # Replace names and roll numbers
        new_text = original_text
        
        # Replace Sahibjit Singh
        if "Sahibjit Singh" in new_text:
            new_text = new_text.replace("Sahibjit Singh (Roll No. 2411789)", "Sehajpreet Singh (Roll No. 2411804)")
            new_text = new_text.replace("Sahibjit Singh", "Sehajpreet Singh")
            new_text = new_text.replace("2411789", "2411804")
            changes_made += 1
        
        # Remove Sargun Singh references
        if "Sargun Singh" in new_text:
            new_text = re.sub(r'\s*and\s*Mr\.\s*Sargun Singh[^.]*?\d{7}\)', '', new_text)
            new_text = re.sub(r'Sargun Singh[^.]*?\d{7}\)', '', new_text)
            new_text = re.sub(r',?\s*Sargun Singh[^\\n]*', '', new_text)
            new_text = new_text.replace("2411799", "")
            changes_made += 1
        
        # Clean up plurals after removing second student
        new_text = new_text.replace("students of", "student of")
        new_text = new_text.replace("have successfully", "has successfully")
        new_text = new_text.replace("they developed", "he developed")
        new_text = new_text.replace("Their performance", "His performance")
        new_text = new_text.replace("We wish them", "We wish him")
        new_text = new_text.replace("their future", "his future")
        
        # Update training dates
        new_text = new_text.replace("4th June 2026 to 10th July 2026", "8th June 2024 to 24th July 2024")
        new_text = new_text.replace("from [20XX-20XX]", "from 2024-2028")
        new_text = new_text.replace("Batch ([20XX-20XX])", "Batch (2024-2028)")
        
        # Update project title
        new_text = new_text.replace("SpringEcom: Enterprise AI-Powered Full-Stack E-Commerce Platform", 
                                    "E-Commerce Website with Admin Authentication System using Spring Boot and React")
        new_text = new_text.replace("Spring Boot 3, Java 17/21, Spring Data JPA, Spring Security (JWT), React 18, Vite, and Google Gemini Generative AI API",
                                    "Spring Boot 3, Java 17, Spring Data JPA, React 18, Vite, PostgreSQL, and modern web technologies")
        
        # Update organization
        new_text = new_text.replace("[Organization Name]", "Udemy (Online Training Platform)")
        new_text = new_text.replace("[Training Manager / In-charge Name]", "Navin Reddy")
        
        # Remove code-like content
        if original_text.strip().startswith("package ") or original_text.strip().startswith("import ") or \
           original_text.strip().startswith("@") or original_text.strip().startswith("public class"):
            # This looks like code - skip it
            para.clear()
            code_sections_removed += 1
            continue
        
        # Apply changes if text was modified
        if new_text != original_text:
            para.clear()
            run = para.add_run(new_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
    
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    original_text = para.text
                    new_text = original_text
                    
                    # Same replacements for tables
                    if "Sahibjit Singh" in new_text:
                        new_text = new_text.replace("Sahibjit Singh (Roll No. 2411789)", "Sehajpreet Singh (Roll No. 2411804)")
                        new_text = new_text.replace("Sahibjit Singh", "Sehajpreet Singh")
                        new_text = new_text.replace("2411789", "2411804")
                    
                    if "Sargun Singh" in new_text:
                        new_text = re.sub(r'Sargun Singh[^\\n]*', '', new_text)
                        new_text = new_text.replace("2411799", "")
                    
                    if new_text != original_text:
                        para.clear()
                        run = para.add_run(new_text.strip())
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(10)
                        changes_made += 1
    
    # Save the edited document
    print(f"Saving: {output_file}")
    doc.save(output_file)
    
    print(f"\n✅ Report edited successfully!")
    print(f"📝 Changes made: {changes_made} text replacements")
    print(f"🗑️  Code sections removed: {code_sections_removed}")
    print(f"📄 Output file: {output_file}")
    print(f"\n👤 Student: Sehajpreet Singh")
    print(f"🎓 Roll No: 2411804")
    print(f"📅 Training: 8 June - 24 July 2024")
    print(f"🏫 Organization: Udemy (Navin Reddy)")

if __name__ == "__main__":
    # Check if file exists
    input_files = [
        "SpringEcom_Project_Report.docx",
        "Project-Report-ECommerce-SpringBoot.docx",
        "Sehajpreet_Singh_2411804_Final_Report.docx"
    ]
    
    input_file = None
    for f in input_files:
        try:
            doc = docx.Document(f)
            input_file = f
            print(f"Found file: {f}")
            break
        except:
            continue
    
    if not input_file:
        print("❌ Could not find the report file!")
        print("Please ensure one of these files exists:")
        for f in input_files:
            print(f"  - {f}")
        sys.exit(1)
    
    output_file = "Sehajpreet_Singh_2411804_Edited_Report.docx"
    edit_report(input_file, output_file)
