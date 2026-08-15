import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml

def add_page_numbers(doc):
    for section in doc.sections:
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        fldChar1 = parse_xml(r'<w:fldChar w:fldCharType="begin" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        instrText = parse_xml(r'<w:instrText xml:space="preserve" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">PAGE</w:instrText>')
        fldChar2 = parse_xml(r'<w:fldChar w:fldCharType="end" xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>')
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)

doc = docx.Document()
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# Title Page
doc.add_paragraph()
p = doc.add_paragraph("AMRITSAR GROUP OF COLLEGES")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(20)
    run.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph("Summer Training Report\nOn")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(14)
    run.font.bold = True

doc.add_paragraph()
p = doc.add_paragraph("E-COMMERCE WEBSITE\nUSING SPRING BOOT")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 73, 125)

doc.add_paragraph()
p = doc.add_paragraph("Submitted in the partial fulfillment of the requirement\nfor the award of degree of")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph("Bachelor of Technology\nin\nCOMPUTER SCIENCE & ENGINEERING\nBatch (2024-2028)")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(14)
    run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph("Submitted by")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(12)
    run.font.bold = True

p = doc.add_paragraph("Sehajpreet Singh\nRoll No: 2411804")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(13)
    run.font.bold = True

doc.add_paragraph()
doc.add_paragraph()
p = doc.add_paragraph("DEPARTMENT OF COMPUTER SCIENCE & ENGINEERING")
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in p.runs:
    run.font.size = Pt(13)
    run.font.bold = True

doc.add_page_break()

add_page_numbers(doc)
doc.save("Sehajpreet_Singh_2411804_Final_Report.docx")
print("✅ Report created: Sehajpreet_Singh_2411804_Final_Report.docx")
